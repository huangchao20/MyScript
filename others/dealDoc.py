# coding=utf-8

import docx
import os
from win32com import client as wc
import re

from app.checkSqlBySh import CheckSql
from app.Mkdirs import Mkdirs
class mainFunc:
    def __init__(self, dpath):
        self.dpath = dpath

    def dealDocxFile(self, filename):
        """
        :function:word文档处理
        :param filename:
        :return:
        """
        print("*******************************************************************************************")
        if os.path.isfile(filename):
            listdir = filename.split('\\')
            listdir.pop()
            if 'XQ' in listdir:
                listdir.pop()
            #组建路径‘F:\\TFS_l\\文档&DB&AFEjar包\\WEEK\\2019\\20190117’
            mkpth = ''
            for p in listdir:
                mkpth = os.path.join(mkpth, p)
            print(mkpth)

            print(filename)
            file = docx.Document(filename)
            #处理docx文档中的表格，处理cfg文件
            tables = file.tables

            for table in tables:
                for i in range(1, len(table.rows)):
                    result = table.cell(1, 0).text
                    if "PORT" in result:
                        print(result)
                mk = Mkdirs(mkpth, result)
                mk.mkNewFile()

            print("段落数: ", str(len(file.paragraphs)))
            # 输出段落编号及段落内容

            for i in range(len(file.paragraphs)):
                # print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
                """
                re.findall(r'\w{2} \w{2}_\d{5}_X_\d{8}.\w{2}', file.paragraphs[i].text)     #匹配fbapDB
                re.findall(r'\w{2} \w{2}_\d{5}_\w{3}_\d{8}_pybak.sh', file.paragraphs[i].text)      #匹配pybak.sh
                re.findall(r'\w{2} \w{2}_\d{5}_\w{4}_\w{2}_\d{8}.sh', file.paragraphs[i].text)        #匹配bsmsDB
                """
                #检查安装手册里面的文档是否存在
                f = filename.split('\\').pop()
                nfilename = filename.replace(f, file.paragraphs[i].text)
                print('新拼接的文件名是[%s]' % nfilename)
                if not os.path.isfile(nfilename):
                    print("安装手册[%s]中的文件[%s]不存在" % (filename, nfilename))
                    return None
                else:   #拼接sbin下面所有的sh、add文件
                    if re.findall(r'\w{2} \w{2}_\d{5}_X_\d{8}.\w{2}', file.paragraphs[i].text):
                        print('fbapdb:[%s]' % file.paragraphs[i].text)
                        cs = CheckSql()
                    elif re.findall(r'\w{2} \w{2}_\d{5}_\w{4}_\w{2}_\d{8}.sh', file.paragraphs[i].text):
                        print('bsmsdb:[%s]' % file.paragraphs[i].text )
                    elif re.findall(r'\w{2} \w{2}_\d{5}_\w{3}_\d{8}_pybak.sh', file.paragraphs[i].text):
                        print('pybak:[%s]' % file.paragraphs[i].text)
                    elif "mkdir" in file.paragraphs[i].text:    #拼接afa、afe的sh脚本
                        mkdirs = file.paragraphs[i].text
                        if "inst1" in mkdirs:
                            print('afa:[%s]' % mkdirs)
                        else:
                            print('afe:[%s]' % mkdirs)
                        mk = Mkdirs(mkpth, mkdirs)
                        mk.mkNewFile()

    def DocxAndDocFile(self, filename):
        if os.path.isfile(filename):
            dfile = os.path.splitext(filename)
            if dfile[1] == ".docx":  # 处理docx文件
                self.dealDocxFile(filename)
            elif dfile[1] == ".doc":  # 处理doc文件
                print("处理doc文件[%s]..." % filename)
                w = wc.Dispatch('Word.Application')
                doc = w.Documents.Open(filename)
                newfile = "".join((dfile[0], ".docx"))
                doc.SaveAs(newfile)
                self.dealDocxFile(newfile)

    def printDocxInfo(self, dpath):
        Keywords = "安装操作"
        filetype = ['.doc', '.docx']
        print(dpath)
        if os.path.isdir(dpath):  # 入参是文件
            for d in os.listdir(dpath):
                filename = os.path.join(dpath, d)
                m = os.path.splitext(d)
                if os.path.isfile(filename) and Keywords in m[0] and m[1] in filetype:
                    self.DocxAndDocFile(filename)
        elif os.path.isfile(dpath):  # 入参是目录
            m = os.path.splitext(dpath)
            if Keywords in m[0] and m[1] in filetype:
                self.DocxAndDocFile(dpath)

    def dirs(self, dpath):
        print("***********************开始处理【%s】*************************" % dpath)
        if not os.path.isdir(dpath):
            print('请检查录入的[%s]目录是否存在！' % dpath)
            return None
        else:
            for dirf in os.listdir(dpath):
                if dirf.startswith(('XQ', 'TR', 'BUG')):
                    ndir = os.path.join(dpath, dirf)
                    self.printDocxInfo(ndir)

"""
from win32com import client as wc 
w = wc.Dispatch('Word.Application') 
# 或者使用下面的方法，使用启动独立的进程： 
# w = wc.DispatchEx('Word.Application') 
doc=w.Documents.Open("E:\\Jupyter\\s.doc") 
doc.SaveAs("E:\\Jupyter\\sa.docx",16)

#读取docx中的文本代码示例
import docx
#获取文档对象
file=docx.Document("D:\\temp\\word.docx")
print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段

#输出每一段的内容
for para in file.paragraphs:
 print(para.text)

#输出段落编号及段落内容
for i in range(len(file.paragraphs)):
 print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)
"""

if __name__ == '__main__':
    dpath = "F:\\TFS_l\\文档&DB&AFEjar包\\WEEK\\2019\\20190117\\XQ-2018-852"
    sp = dpath.split('\\').pop()
    mainfunc = mainFunc(dpath)

    if sp.startwith('XQ'):
        mainfunc.printDocxInfo(dpath)
    else:
        mainfunc.dirs(dpath)