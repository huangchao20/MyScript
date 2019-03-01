#encoding=utf-8
import os
import docx
from win32com import client as wc
import re

from app.checkSqlBySh import CheckSql
from app.Mkdirs import Mkdirs

class DealPack:
    def __init__(self, dpath):
        self.dpath = dpath

    def GetShName(self, str1, filename, docxname, xqdir):
        """
        :param str1: sbin脚本添加的内容
        :param filename: docx的绝对路径文件名
        :param docxname:docx的文件名
        :param xqdir:XQ号
        :return:
        """
        tli = str1.split(" ")
        for name in tli:
            if name.endswith('.sh'):
                shname = filename.replace(docxname, name)
                t = str1.split(' ')
                for s in t:
                    if 'pybak' in s or 'pyback' in s:
                        mkdirs = "/".join((xqdir, s))
                    else:
                        s = " ".join((xqdir, s))
                        mkdirs = ' '.join(('install ', s))
        #shname格式：'E:\SVN\2019\20190110w\特色业务平台\t6\fbap.20190110rw.t6\XQ-2018-801\TS_77044_AFA_20100117_pybak.sh'
        return shname, mkdirs

    def __dealDocxFile(self, filename):
        """
         :function:word文档处理
         :param filename:
         :return:
         """
        print("【*****__dealDocxFile开始处理文档[%s]*****】" % filename)
        if os.path.isfile(filename):
            sfile = ('XQ', 'BUG', 'TR')
            listdir = filename.split('\\')
            listdir.pop()
            if listdir[-1].startswith(sfile):
                xqdir = listdir.pop()
            # 组建路径‘F:\\TFS_l\\文档&DB&AFEjar包\\WEEK\\2019\\20190117’
            mkpth = '\\'.join((listdir[0], ""))
            for p in listdir[1:]:
                print(p)
                mkpth = os.path.join(mkpth, p)
            print('sbin的目录是[%s], XQ号的名称是:[%s]' % (mkpth, xqdir))
            print(filename)
            file = docx.Document(filename)
            # 处理docx文档中的表格，处理cfg文件
            tables = file.tables

            for table in tables:
                for i in range(1, len(table.rows)):
                    result = table.cell(1, 0).text
                    if "PORT" in result:
                        print(result)
                mk = Mkdirs(mkpth, result)

            print("段落数: ", str(len(file.paragraphs)))
            # 输出段落编号及段落内容
            for i in range(len(file.paragraphs)):
                # print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
                """
                re.findall(r'\w{2} \w{2}_\d{5}_X_\d{8}.\w{2}', file.paragraphs[i].text)     #匹配fbapDB
                re.findall(r'\w{2} \w{2}_\d{5}_\w{3}_\d{8}_pybak.sh', file.paragraphs[i].text)      #匹配pybak.sh
                re.findall(r'\w{2} \w{2}_\d{5}_\w{4}_\w{2}_\d{8}.sh', file.paragraphs[i].text)        #匹配bsmsDB
                """
                # 检查安装手册里面的文档是否存在
                fp = file.paragraphs[i].text.replace('\t', '')
                fp.rstrip()
                if "mkdir" in fp:  # 拼接afa、afe的sh脚本
                    print('>>>>>>>>>>>>>>>>>>>开始拼接afa、afe脚本<<<<<<<<<<<<<<<<<<<<')
                    mkdirs = fp
                    if "inst1" in mkdirs:
                        print('afa:[%s]' % mkdirs)
                    else:
                        print('afe:[%s]' % mkdirs)
                    mk = Mkdirs(mkpth, mkdirs)
                elif fp.endswith('.sh'):
                    print('>>>>>>>>>>>>>>>>>>>>开始处理脚本[%s]<<<<<<<<<<<<<<<<<<<<<' % fp)
                    t = fp.split(' ')
                    for n in range(len(t)):
                        if t[n].endswith('.sh'):
                            rename = t[n]
                    f = filename.split('\\').pop()
                    # nfilename = filename.replace(f, file.paragraphs[i].text)
                    nfilename = filename.replace(f, rename)
                    print('新拼接的文件名是[%s]' % nfilename)
                    if not os.path.isfile(nfilename):
                        print("安装手册[%s]中的文件[%s]不存在" % (filename, nfilename))
                        return None
                    else:  # 拼接sbin下面所有的sh、add文件
                        if ".sh" in fp:
                            for s in fp.split(' '):
                                if s.endswith('.sh'):
                                    shs = s
                        if re.findall(r'\w{2}_\d{5}_\w_\d{8}.\w{2}', shs):
                            print('【************fbapDB开始处理：[%s]************】' %fp)
                            shn, mkdirs = self.GetShName(fp, filename, f, xqdir)
                        elif re.findall(r'\w{2}_\d{5}_\w{4}_\w{2}_\d{8}.sh', shs):
                            print('【************bsmsDB开始处理：[%s]************】' % fp)
                            print('bsmsdb:[%s]' % file.paragraphs[i].text)
                            shn, mkdirs = self.GetShName(fp, filename, f, xqdir)
                        elif re.findall(r'\w{2}_\d{5}_\w{3}_\d{8}_\w{5,6}.sh', shs):
                            print('【************pybak开始处理：[%s]************】' % fp)
                            # mkdirs = '/'.join((xqdir, file.paragraphs[i].text))
                            shn, mkdirs = self.GetShName(fp, filename, f, xqdir)
                            Mkdirs(mkpth, mkdirs)
                        else:
                            continue
                        print(shn)
                        CheckSql(shn, mkpth, mkdirs)

    def TextFile(self):
        dlist = self.dpath.split('\\')
        print(dlist)
        filename = self.dpath
        if os.path.isfile(self.dpath):
            print("DealPack-->>self.dpath[%s]" % filename)
            dfile = os.path.splitext(filename)
            if dfile[1] == ".docx":  # 处理docx文件
                self.__dealDocxFile(filename)
            elif dfile[1] == ".doc":  # 将doc文件转换成docx文件
                print("处理doc文件[%s]..." % filename)
                w = wc.Dispatch('Word.Application')
                doc = w.Documents.Open(filename)
                newfile = "".join((dfile[0], ".docx"))
                doc.SaveAs(newfile)
                self.__dealDocxFile(newfile)
        return '000000'

if __name__ == '__main__':
    dpath = 'F:\\黄小宝的宝\\测试目录\\fbap.20190110rw.t6\\XQ-2018-801\\TS_75994_安装操作部署手册.docx'
    t = DealPack(dpath)
    t.TextFile()